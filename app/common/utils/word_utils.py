"""Word utility functions (python-docx)."""

from __future__ import annotations

import io

from app.core.result import AppError, Err, Ok, Result

# Word テーブル解析の戻り値型
# 各テーブルは「キー・値ペアの行リスト」として表現される
# 例: [{"key": "製品名", "value": "XYZ"}, {"key": "申請番号", "value": "2026-001"}]
WordTableRow = dict[str, str]
WordTableList = list[list[WordTableRow]]


def parse_word_tables(
    file_content: bytes,
) -> Result[WordTableList]:
    """Word ファイル（.docx）の全テーブルを解析して返す。

    Word テンプレート構造を前提:
      - 2 列テーブル（項目名 │ 内容）
      - 複数テーブルが存在する場合はすべて返す

    DB への書き込みは行わない（プレビュー・データ取得専用）。

    【初学者向け】
    - python-docx の Document クラスでファイルを読み込む。
    - doc.tables で全テーブルを取得し、row.cells でセルを取得する。
    - key が空行・ブランク行はスキップしてきれいなデータのみ返す。

    Args:
        file_content: .docx ファイルの生バイト列。

    Returns:
        Ok(tables) を返す。
        - tables: テーブルのリスト。各テーブルは {"key": str, "value": str} の辞書リスト。
        失敗時は Err(VALIDATION) を返す。

    Example::

        result = parse_word_tables(file_content)
        if is_err(result):
            return result
        for table in result.value:
            for row in table:
                print(f"{row['key']}: {row['value']}")
    """
    try:
        from docx import Document  # type: ignore[import-untyped]

        doc = Document(io.BytesIO(file_content))
        tables: WordTableList = []

        for table in doc.tables:
            rows: list[WordTableRow] = []
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    key = cells[0].text.strip()
                    value = cells[1].text.strip()
                    # ヘッダー行・ブランク行はスキップ
                    if key:
                        rows.append({"key": key, "value": value})
            if rows:
                tables.append(rows)

    except Exception as exc:
        return Err(
            error=AppError(
                type="VALIDATION",
                message="Failed to parse Word file. Make sure it is a valid .docx file.",
                details=str(exc),
            )
        )

    return Ok(value=tables)
